from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUT_PATH = Path("out_candidates") / "simulation_component_section.docx"


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, sub=False, sup=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.subscript = sub
    run.font.superscript = sup


def add_text_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True)
    return p


def add_equation(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    for part in parts:
        run = p.add_run(part["text"])
        set_font(
            run,
            name="Cambria Math",
            size=12,
            italic=part.get("italic", False),
            sub=part.get("sub", False),
            sup=part.get("sup", False),
        )
    return p


def t(text, **kwargs):
    item = {"text": text}
    item.update(kwargs)
    return item


def main():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    add_heading(doc, "Simulation Component", 1)
    add_text_paragraph(
        doc,
        "The simulation component was designed as a configuration-driven discrete-event "
        "simulation framework for microservice-based systems. A simulation run is defined "
        "through a scenario-based YAML specification that describes the available hosts, "
        "deployed services, replica counts, endpoint-level CPU and memory demands, downstream "
        "service dependencies, latency characteristics, workload arrival patterns, and optional "
        "runtime policies. At execution time, the scenario is parsed into an internal state "
        "composed of an event engine, resource manager, interaction manager, policy manager, "
        "workload generator, and metrics collector. This structure enables the simulator to "
        "reproduce not only service execution behavior, but also deployment-level resource "
        "constraints and runtime control decisions within a unified experimental environment.",
    )

    add_heading(doc, "Event-Driven Execution Model", 2)
    add_text_paragraph(
        doc,
        "The simulator follows a discrete-event execution model in which all state transitions "
        "are represented as timestamped events stored in a priority queue. Events are processed "
        "in non-decreasing order of simulation time, and event priority is used as a tie-breaking "
        "rule when two events share the same timestamp. This allows the engine to advance directly "
        "from one meaningful system change to the next, rather than iterating through fixed time "
        "steps. The main event types include request arrival, request start, request completion, "
        "downstream invocation, drain processing for scaling operations, and simulation termination. "
        "As a result, the execution model is efficient while still preserving causal ordering "
        "across concurrent service interactions.",
    )

    add_heading(doc, "Workload and Service-Time Models", 2)
    add_text_paragraph(
        doc,
        "Incoming workload is generated through several arrival processes, including Poisson, "
        "constant, Gaussian, uniform, and burst-oriented patterns. For Poisson traffic, the "
        "inter-arrival time is sampled from an exponential distribution with rate λ:",
    )
    add_equation(
        doc,
        [
            t("Δt ∼ Exp(λ), λ = rate"),
            t("rps", sub=True),
            t("."),
        ],
    )
    add_text_paragraph(doc, "For deterministic traffic, the simulator uses:")
    add_equation(
        doc,
        [
            t("Δt = 1 / rate"),
            t("rps", sub=True),
            t("."),
        ],
    )
    add_text_paragraph(
        doc,
        "For Gaussian arrival behavior, the configured request-rate parameters are transformed "
        "into inter-arrival parameters according to:",
    )
    add_equation(
        doc,
        [
            t("μ", italic=True),
            t("Δt", sub=True),
            t(" = 1 / μ", italic=True),
            t("rps", sub=True),
            t(", σ", italic=True),
            t("Δt", sub=True),
            t(" = σ", italic=True),
            t("rps", sub=True),
            t(" / μ", italic=True),
            t("rps", sub=True),
            t("2", sup=True),
            t(","),
        ],
    )
    add_text_paragraph(doc, "after which inter-arrival times are sampled from:")
    add_equation(
        doc,
        [
            t("Δt ∼ N(μ", italic=True),
            t("Δt", sub=True),
            t(", σ", italic=True),
            t("Δt", sub=True),
            t(")."),
        ],
    )
    add_text_paragraph(
        doc,
        "Endpoint execution time is modeled as the sum of sampled CPU time, sampled network "
        "latency, and an estimated queueing delay. CPU demand is generated as:",
    )
    add_equation(
        doc,
        [
            t("T"),
            t("cpu", sub=True),
            t(" ∼ N(μ", italic=True),
            t("cpu", sub=True),
            t(", σ", italic=True),
            t("cpu", sub=True),
            t("),"),
        ],
    )
    add_text_paragraph(doc, "while endpoint network delay is modeled by:")
    add_equation(
        doc,
        [
            t("T"),
            t("net", sub=True),
            t(" ∼ N(μ", italic=True),
            t("net", sub=True),
            t(", σ", italic=True),
            t("net", sub=True),
            t("),"),
        ],
    )
    add_text_paragraph(
        doc,
        "with negative sampled values clamped to zero. Queueing delay is approximated from the "
        "current queue length and the expected mean service demand:",
    )
    add_equation(
        doc,
        [
            t("T"),
            t("queue", sub=True),
            t(" = L"),
            t("q", sub=True),
            t(" ⋅ (μ", italic=True),
            t("cpu", sub=True),
            t(" + μ", italic=True),
            t("net", sub=True),
            t("),"),
        ],
    )
    add_text_paragraph(
        doc,
        "where Lq denotes the queue length at the selected instance. The total request processing "
        "time is therefore estimated as:",
    )
    add_equation(
        doc,
        [
            t("T"),
            t("proc", sub=True),
            t(" = T"),
            t("cpu", sub=True),
            t(" + T"),
            t("net", sub=True),
            t(" + T"),
            t("queue", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(doc, "and the observed end-to-end request latency is computed as:")
    add_equation(
        doc,
        [
            t("T"),
            t("latency", sub=True),
            t(" = t"),
            t("completion", sub=True),
            t(" − t"),
            t("arrival", sub=True),
            t("."),
        ],
    )
    add_text_paragraph(
        doc,
        "For service-to-service dependencies, downstream invocations are generated after request "
        "completion and can also include an additional sampled call latency. When multiple "
        "downstream calls are possible, the simulator supports branching behavior through "
        "stochastic rounding of the configured mean call count or, alternatively, probability-based "
        "call selection.",
    )

    add_heading(doc, "Resource and Routing Models", 2)
    add_text_paragraph(
        doc,
        "The resource model operates at both service-instance and host levels. Requests are routed "
        "to service replicas using round-robin selection, thereby distributing traffic across "
        "routable instances in a simple and reproducible manner. Each service instance tracks CPU "
        "consumption over a sliding observation window and computes CPU utilization as:",
    )
    add_equation(
        doc,
        [
            t("U"),
            t("cpu", sub=True),
            t(" = (T"),
            t("cpu", sub=True),
            t("window", sup=True),
            t(" / T"),
            t("window", sub=True),
            t(") / C"),
            t("cpu", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(
        doc,
        "where Tcpuwindow is the CPU time accumulated in the current window, Twindow is the "
        "window duration, and Ccpu is the CPU capacity allocated to the instance. Memory "
        "utilization is calculated as:",
    )
    add_equation(
        doc,
        [
            t("U"),
            t("mem", sub=True),
            t(" = M"),
            t("active", sub=True),
            t(" / M"),
            t("allocated", sub=True),
            t("."),
        ],
    )
    add_text_paragraph(
        doc,
        "At the host level, utilization is aggregated from all resident instances. Host CPU "
        "utilization is computed as:",
    )
    add_equation(
        doc,
        [
            t("U"),
            t("host", sub=True),
            t("cpu", sup=True),
            t(" = Σ"),
            t("i", sub=True),
            t(" U"),
            t("i", sub=True),
            t("cpu", sup=True),
            t(" C"),
            t("i", sub=True),
            t(" / C"),
            t("host", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(doc, "and host memory utilization is computed as:")
    add_equation(
        doc,
        [
            t("U"),
            t("host", sub=True),
            t("mem", sup=True),
            t(" = Σ"),
            t("i", sub=True),
            t(" M"),
            t("i", sub=True),
            t("active", sup=True),
            t(" / M"),
            t("host", sub=True),
            t("."),
        ],
    )
    add_text_paragraph(
        doc,
        "These formulations allow the simulator to capture both local service pressure and global "
        "infrastructure saturation during runtime.",
    )

    add_heading(doc, "Runtime Control and Adaptive Behavior", 2)
    add_text_paragraph(
        doc,
        "The simulation component also includes runtime control algorithms that enable "
        "experimentation with operational policies. Request admission can be restricted using a "
        "token-bucket rate limiter, where token replenishment after elapsed time Δt is given by:",
    )
    add_equation(
        doc,
        [
            t("Δtokens = ⌊Δt ⋅ r⌋,"),
        ],
    )
    add_text_paragraph(
        doc,
        "with r representing the refill rate. Circuit breaking is modeled through the standard "
        "closed-open-half-open state machine using configurable failure and recovery thresholds. "
        "Configurable retry behavior is represented through backoff functions, including the "
        "exponential form:",
    )
    add_equation(
        doc,
        [
            t("B(a) = B"),
            t("0", sub=True),
            t(" ⋅ 2"),
            t("a−1", sup=True),
            t(","),
        ],
    )
    add_text_paragraph(
        doc,
        "where B0 is the base backoff duration and a is the retry attempt index. In addition, "
        "autoscaling-related control logic is supported through threshold-based scaling decisions "
        "with hysteresis. Scale-up is triggered when measured utilization exceeds a target "
        "threshold:",
    )
    add_equation(
        doc,
        [
            t("U"),
            t("cpu", sub=True),
            t(" > U"),
            t("target", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(doc, "while scale-down is permitted only when utilization falls below a reduced threshold:")
    add_equation(
        doc,
        [
            t("U"),
            t("cpu", sub=True),
            t(" < αU"),
            t("target", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(
        doc,
        "where α = 0.8 is the hysteresis factor used to reduce oscillatory scaling.",
    )

    add_heading(doc, "Metric Aggregation, Interfaces, and Experiment Support", 2)
    add_text_paragraph(
        doc,
        "During execution, the simulator records time-series observations for request count, "
        "errors, latency, CPU utilization, memory utilization, queue length, throughput, and "
        "concurrency. These samples are aggregated into run-level performance summaries that "
        "include the arithmetic mean together with percentile-based latency measures such as P50, "
        "P95, and P99. Throughput is calculated as:",
    )
    add_equation(
        doc,
        [
            t("Throughput = N"),
            t("requests", sub=True),
            t(" / T"),
            t("sim", sub=True),
            t(","),
        ],
    )
    add_text_paragraph(
        doc,
        "where Nrequests is the total number of processed requests and Tsim is the simulation "
        "duration. Percentiles are derived from sorted sample sets using interpolation, allowing "
        "robust summary statistics for skewed latency distributions.",
    )
    add_text_paragraph(
        doc,
        "Beyond the mathematical models and runtime algorithms, the simulation component also "
        "includes scenario parsing, run lifecycle orchestration, dynamic mid-run reconfiguration "
        "of workload and resource settings, and HTTP/gRPC-based exposure of run state and metrics. "
        "Consequently, the component serves not only as an execution engine for microservice "
        "simulation, but also as an experimental platform for controlled what-if analysis, live "
        "monitoring, and integration with higher-level optimization and orchestration workflows.",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    main()
